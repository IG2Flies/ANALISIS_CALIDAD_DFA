import streamlit as st
import google.generativeai as genai
import pandas as pd
import tempfile
import os
import time
import json
import zipfile
import re
import unicodedata
from io import BytesIO

# --- LIBRERÍAS TÉCNICAS ---
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import mutagen 
except ImportError:
    st.error("⚠️ Faltan librerías. Ejecuta: python -m pip install python-docx mutagen openpyxl pandas")
    st.stop()

st.set_page_config(page_title="Dashboard DFA Pro", layout="wide", page_icon="📊")
st.title("📊 Auditoría Integral - Organización por Servicios")

# --- BASE DE DATOS DE OPERADORES Y SERVICIOS (Transcipción de tu imagen) ---
# Esto permite asignar automáticamente el servicio según el nombre del operador
DB_OPERADORES = {
    "FAMILIAS": ["Cristina Ruiz", "Elena Pinilla"],
    "TUTELAS": ["Marta Rodriguez"],
    "DIR. GEN. FAMILIAS": ["Carmen Cuadrado", "Andrea Sanz", "Elena Canera", "Javier García"],
    "IASS": ["Inma Marcen", "María José Cabrero", "Sena Echegoyen", "Teresa Alfaro", "Ana Isabel Mojares", "Daniel Aznar", "María José Monterde", "Rosana Soriano", "Willy Putze"],
    "AT. DISCAPACIDAD": ["Marta Candado"],
    "AYTO ZARAGOZA": ["Yanila Marco", "Asunción Fernández", "Assia Hanib", "Carmen Casas", "Marina Martín", "Cristina Arruga", "Mario González", "Silvia Dobre", "Pilar Alcubierre", "Laura Nuñez", "Mayte Moreno", "Noelia Navarro", "Ana Crespo", "Sara Galvez", "Olena Kostenko", "Andrea Soria Zapata"],
    "ATAM": ["Asunción Pérez"],
    "AT. DFA": ["Lorena Gracia", "Raul Mercado"],
    # Servicios sin lista de operadores fija (se asignarán si Gemini lo detecta o manualmente)
    "ALHAMBRA": [],
    "TELF DEL MAYOR": [],
    "TRESUELVE": [],
    "SCB": [],
    "MUM": [],
    "ESTUDIOS DE MERCADO": [],
    "MÉMORA": []
}

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("1. Configuración")
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
        st.success("🔑 Llave cargada")
    else:
        api_key = st.text_input("Tu API Key", type="password")
    
    st.header("2. Modelo IA")
    modelos = ["models/gemini-3-pro-preview", "models/gemini-1.5-pro", "models/gemini-2.5-flash"]
    modelo_elegido = st.selectbox("Modelo:", modelos, index=0)
    
    st.divider()
    st.info("📂 El sistema organizará las carpetas por SERVICIO > OPERADOR automáticamente.")

# --- FUNCIONES AUXILIARES ---
def normalizar_texto(texto):
    """Quita tildes y pone minúsculas para comparar nombres."""
    if not texto: return ""
    texto = str(texto).lower()
    return ''.join(c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn')

def detectar_servicio(nombre_operador_detectado):
    """Busca en la DB a qué servicio pertenece el operador."""
    nombre_limpio = normalizar_texto(nombre_operador_detectado)
    
    for servicio, lista_nombres in DB_OPERADORES.items():
        for nombre_db in lista_nombres:
            # Si el nombre de la DB está contenido en lo que detectó la IA (ej: "Sara Galvez" en "Soy Sara Galvez")
            if normalizar_texto(nombre_db) in nombre_limpio:
                return servicio
    
    return "OTROS SERVICIOS" # Si no encuentra coincidencia

def limpiar_nombre(texto):
    if not texto: return "Desconocido"
    return re.sub(r'[\\/*?:"<>|]', "", str(texto)).strip()

def obtener_metadatos_forenses(ruta_archivo):
    info_dict = {"Tamaño": f"{os.path.getsize(ruta_archivo) / 1024:.2f} KB"}
    try:
        audio = mutagen.File(ruta_archivo)
        if audio and audio.info:
            secs = int(audio.info.length)
            info_dict["Duracion_Texto"] = f"{secs // 60} min {secs % 60} seg"
            if hasattr(audio.info, 'bitrate'): 
                info_dict["Bitrate"] = f"{audio.info.bitrate / 1000:.0f} kbps"
    except: pass
    return info_dict

def crear_docx_individual(data, servicio_asignado):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    nota = data.get("5_evaluacion_detallada", {}).get("calificacion_general", "?")
    op = data.get("1_identificacion", {}).get("operador", "N/A")
    titulo = doc.add_heading(f"[{servicio_asignado}] {op} - Nota: {nota}", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. FICHA TÉCNICA
    doc.add_heading("1. FICHA TÉCNICA", level=1)
    p = doc.add_paragraph()
    p.add_run("SERVICIO: ").bold = True
    p.add_run(servicio_asignado).font.color.rgb = RGBColor(200, 0, 0)
    
    p = doc.add_paragraph()
    p.add_run("OPERADOR: ").bold = True
    p.add_run(op).font.color.rgb = RGBColor(0, 51, 102)

    ident = data.get("1_identificacion", {})
    # Tabla
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for k, v in ident.get("metadatos_tecnicos", {}).items():
        row = table.add_row().cells
        row[0].text = k.replace("_", " ").upper()
        row[1].text = str(v)

    # 2. ANÁLISIS
    doc.add_heading("2. ANÁLISIS FORENSE", level=1)
    doc.add_paragraph(data.get("3_resumen_ejecutivo", {}).get("cronologia", "-"))
    
    concl = data.get("4_conclusiones", {})
    p = doc.add_paragraph()
    p.add_run("🔴 Puntos Dolor: ").bold = True
    p.add_run(concl.get("puntos_dolor", "-")).font.color.rgb = RGBColor(180, 0, 0)
    
    # 3. FEEDBACK
    doc.add_heading("3. COACHING", level=1)
    doc.add_paragraph(data.get("7_coaching", {}).get("mensaje_directo", "-")).italic = True
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_zip_dashboard(resultados):
    zip_buffer = BytesIO()
    filas_excel = []

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for res in resultados:
            op = limpiar_nombre(res.get("1_identificacion", {}).get("operador", "Desconocido"))
            
            # --- LÓGICA DE ORGANIZACIÓN ---
            # 1. Detectamos el servicio basado en el nombre del operador
            servicio = detectar_servicio(op)
            
            # 2. Si es desconocido, miramos si la IA detectó algún servicio en el audio
            if servicio == "OTROS SERVICIOS":
                servicio_ia = res.get("2_asunto", {}).get("servicio_detectado", "")
                if servicio_ia: servicio = limpiar_nombre(servicio_ia)

            nota = res.get("5_evaluacion_detallada", {}).get("calificacion_general", "0")
            titulo = limpiar_nombre(res.get("titulo_informe", "doc"))
            
            # RUTA: SERVICIO / OPERADOR / [NOTA] ARCHIVO
            nombre_docx = f"[{nota}] {titulo}.docx"
            ruta_zip = f"{servicio}/{op}/{nombre_docx}"
            
            # Añadir Word al ZIP
            zf.writestr(ruta_zip, crear_docx_individual(res, servicio).getvalue())
            
            # Excel Link
            link = f'=HYPERLINK("{ruta_zip}", "Abrir Informe")'
            
            filas_excel.append({
                "Servicio": servicio, # Nueva Columna
                "Operador": op,
                "Nota": float(nota) if str(nota).replace('.','',1).isdigit() else 0,
                "Enlace": link,
                "Archivo": titulo,
                "Duración": res.get("1_identificacion", {}).get("metadatos_tecnicos", {}).get("Duracion_Texto", "-"),
                "Ring Time": res.get("1_identificacion", {}).get("metadatos_tecnicos", {}).get("Ring_Time_Estimado", "-"),
                "Colgó": res.get("1_identificacion", {}).get("metadatos_tecnicos", {}).get("Quien_Cuelga", "-")
            })
            
        # Generar Excel Maestro
        if filas_excel:
            df = pd.DataFrame(filas_excel)
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name="Dashboard Global")
            zf.writestr("00_CUADRO_MANDO_SERVICIOS.xlsx", excel_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer

def analizar_con_ia(audio, meta_txt, key, modelo, prompt):
    genai.configure(api_key=key)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tfile:
        tfile.write(audio.read())
        path = tfile.name

    try:
        meta_py = obtener_metadatos_forenses(path)
        full_prompt = f"{prompt}\nDATOS TÉCNICOS: {meta_py}\nMETADATOS: {meta_txt}"
        
        myfile = genai.upload_file(path)
        while myfile.state.name == "PROCESSING": time.sleep(1)
        
        model = genai.GenerativeModel(modelo)
        resp = model.generate_content([myfile, full_prompt])
        myfile.delete()
        return json.loads(resp.text.replace("```json", "").replace("```", "").strip())
    except Exception as e:
        return {"Error": str(e), "titulo_informe": audio.name}
    finally:
        if os.path.exists(path): os.unlink(path)

# --- PROMPT ACTUALIZADO ---
PROMPT_MAESTRO = """
ACTÚA COMO: Auditor de Calidad.
OBJETIVO: Identificar operador y servicio, y evaluar calidad.

INSTRUCCIONES CLAVE:
1. IDENTIFICACIÓN: Extrae el nombre del operador si se presenta (ej: "Soy Sara Galvez").
2. SERVICIO: Intenta deducir el servicio por el contexto (ej: "Ayuntamiento", "Telesasistencia", "Mémora", "IASS").
3. TIEMPOS: Estima Ring Time y quién cuelga.

OUTPUT JSON:
{
    "titulo_informe": "Nombre archivo",
    "1_identificacion": {
        "operador": "Nombre o 'Desconocido'",
        "metadatos_tecnicos": { "Duracion_Texto": "...", "Ring_Time_Estimado": "...", "Quien_Cuelga": "..." }
    },
    "2_asunto": {
        "motivo": "...",
        "servicio_detectado": "Ej: Ayto Zaragoza / IASS / Alhambra (si se menciona)"
    },
    "3_resumen_ejecutivo": { "cronologia": "..." },
    "4_conclusiones": { "puntos_dolor": "..." },
    "5_evaluacion_detallada": { "calificacion_general": "8.5", "analisis_tiempos": { "justificacion": "..." } },
    "7_coaching": { "mensaje_directo": "..." }
}
"""

# --- INTERFAZ PRINCIPAL ---
uploaded = st.file_uploader("📂 Sube Audios (.mp3) y Metadatos", accept_multiple_files=True)

if uploaded and st.button("🚀 AUDITAR ORGANIZADO POR SERVICIOS"):
    audios = [f for f in uploaded if f.name.lower().endswith('.mp3')]
    extras = {os.path.splitext(f.name)[0]: f.getvalue().decode('utf-8', errors='ignore') for f in uploaded if f.name.endswith(('.xml','.txt'))}

    if not api_key:
        st.error("Falta API Key")
    else:
        # 1. PROCESAMIENTO
        bar = st.progress(0)
        resultados = []
        status_box = st.empty()
        
        for i, audio in enumerate(audios):
            status_box.info(f"Analizando {i+1}/{len(audios)}: {audio.name}")
            meta = extras.get(os.path.splitext(audio.name)[0], "")
            res = analizar_con_ia(audio, meta, api_key, modelo_elegido, PROMPT_MAESTRO)
            if "titulo_informe" not in res: res["titulo_informe"] = audio.name
            resultados.append(res)
            bar.progress((i+1)/len(audios))
            if "pro" in modelo_elegido: time.sleep(3)
        
        status_box.success("✅ Análisis completado.")
        
        # 2. VISUALIZACIÓN DASHBOARD (NUEVO: POR SERVICIO)
        st.divider()
        st.subheader("📈 Estadísticas por Servicio")
        
        datos_grafica = []
        for r in resultados:
            op = r.get("1_identificacion", {}).get("operador", "Desconocido")
            # Usamos la misma lógica de detección que para el ZIP
            serv = detectar_servicio(op)
            if serv == "OTROS SERVICIOS":
                s_ia = r.get("2_asunto", {}).get("servicio_detectado", "")
                if s_ia: serv = limpiar_nombre(s_ia)
            
            nota = r.get("5_evaluacion_detallada", {}).get("calificacion_general", 0)
            try: nota = float(nota)
            except: nota = 0
            datos_grafica.append({"Servicio": serv, "Operador": op, "Nota": nota})
        
        df_dash = pd.DataFrame(datos_grafica)
        
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Nota Media por Servicio**")
            # Agrupar por Servicio y sacar media
            if not df_dash.empty:
                st.bar_chart(df_dash.groupby("Servicio")["Nota"].mean())
            else:
                st.write("Sin datos.")
        
        with col2:
            st.write("**Detalle Operadores**")
            st.dataframe(df_dash)

        # 3. DESCARGA ZIP
        zip_final = generar_zip_dashboard(resultados)
        st.download_button(
            label="📦 DESCARGAR PACK SERVICIOS (.ZIP)",
            data=zip_final,
            file_name="Auditoria_Servicios_DFA.zip",
            mime="application/zip",
            type="primary"
        )